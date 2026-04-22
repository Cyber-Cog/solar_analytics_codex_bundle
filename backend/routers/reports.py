"""
backend/routers/reports.py
==========================
Reporting module — generates plant performance reports in **PDF, XLSX, and DOCX**
formats with KPIs, inverter tables, loss-category summaries and matplotlib charts.

All three formats share the same underlying `_gather_report_payload()` so numbers
match across outputs.  Optional dependencies (reportlab, python-docx, matplotlib)
are imported lazily so the API still boots when a format is unavailable —
clients receive an HTTP 503 with a clear message instead of a 500 stack trace.

Endpoints
---------
GET  /api/reports/options           → which formats are actually installable
POST /api/reports/generate          → streams a generated document
"""

from __future__ import annotations

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import text
from pydantic import BaseModel, Field
from typing import Any, Dict, List, Optional
from datetime import datetime, date
from io import BytesIO
import base64
import logging
import os

from database import get_db
from models import Plant, User
from auth.routes import get_current_user

logger = logging.getLogger("reports")
router = APIRouter(prefix="/api/reports", tags=["Reports"])


# ──────────────────────────────────────────────────────────────────────────────
# Optional-dependency shims
# ──────────────────────────────────────────────────────────────────────────────
def _have(mod: str) -> bool:
    """Quietly probe whether a module is importable without raising on failure."""
    try:
        __import__(mod)
        return True
    except Exception:
        return False


_CAPS = {
    "pdf": _have("reportlab"),
    "xlsx": _have("openpyxl"),
    "docx": _have("docx"),
    "html": True,  # HTML is pure-python (stdlib) — always available
    "charts": _have("matplotlib"),
}


# ──────────────────────────────────────────────────────────────────────────────
# Request model
# ──────────────────────────────────────────────────────────────────────────────
class ReportRequest(BaseModel):
    plant_id: str
    date_from: Optional[str] = None
    date_to: Optional[str] = None
    format: str = Field("pdf", pattern="^(pdf|xlsx|docx|html)$")
    sections: List[str] = Field(
        default_factory=lambda: ["overview", "kpis", "inverters", "losses", "faults", "energy_trend"]
    )
    title: Optional[str] = None


# ──────────────────────────────────────────────────────────────────────────────
# Data gathering — reuses dashboard/fault helpers so numbers match the UI.
# ──────────────────────────────────────────────────────────────────────────────
def _gather_report_payload(
    db: Session, plant_id: str, date_from: Optional[str], date_to: Optional[str]
) -> Dict[str, Any]:
    """Pull every dataset the report can show — station, KPIs, inverters, energy trend, losses."""
    from routers.dashboard import dashboard_bundle as _bundle_endpoint  # type: ignore

    # Fake a minimal `current_user` that the endpoint signature wants. The bundle only
    # needs db + plant_id + dates, so a sentinel user object is safe.
    class _U:  # minimal stand-in; downstream only reads attrs we don't reference.
        id = 0
        username = "__report__"

    try:
        bundle = _bundle_endpoint(  # type: ignore[arg-type]
            plant_id=plant_id,
            date_from=date_from,
            date_to=date_to,
            db=db,
            current_user=_U(),
        )
    except Exception as exc:
        logger.exception("dashboard_bundle failed")
        raise HTTPException(500, f"Dashboard data unavailable: {exc}")

    losses: Dict[str, Any] = {}
    try:
        from routers.faults import _unified_fault_categories_core  # type: ignore

        core = _unified_fault_categories_core(db, plant_id, date_from, date_to, _U())
        losses = {"categories": core.get("categories", []), "totals": core.get("totals", {})}
    except Exception:
        logger.warning("unified fault categories unavailable", exc_info=True)

    return {"bundle": bundle, "losses": losses}


# ──────────────────────────────────────────────────────────────────────────────
# Chart rendering (matplotlib) — returns PNG bytes
# ──────────────────────────────────────────────────────────────────────────────
def _render_energy_chart_png(energy_points: List[Dict[str, Any]]) -> Optional[bytes]:
    if not _CAPS["charts"] or not energy_points:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        labels = [str(p.get("date") or p.get("day") or "") for p in energy_points]
        # Dashboard bundle emits `actual_kwh`; fall back to `energy_kwh`/`value` for safety.
        values = [
            float(p.get("actual_kwh") or p.get("energy_kwh") or p.get("value") or 0)
            for p in energy_points
        ]
        targets = [
            float(p.get("target_kwh") or p.get("target") or 0) for p in energy_points
        ]
        has_target = any(t > 0 for t in targets)

        fig, ax = plt.subplots(figsize=(9.5, 4.0), dpi=180)
        fig.patch.set_facecolor("white")
        bars = ax.bar(range(len(values)), values, color="#2563eb", width=0.66,
                      edgecolor="#1e40af", linewidth=0.5, label="Actual")
        if has_target:
            ax.plot(range(len(targets)), targets, color="#f97316", marker="o",
                    linewidth=2.0, markersize=5, label="Target",
                    markeredgecolor="white", markeredgewidth=1.2, zorder=5)
            ax.legend(loc="upper right", fontsize=10, frameon=True,
                      facecolor="white", edgecolor="#cbd5e1")
        ax.set_ylabel("Energy (kWh)", fontsize=11, weight="bold", color="#1e293b")
        ax.set_title("Daily Energy Generation — Actual vs Target",
                     fontsize=13, weight="bold", pad=12, color="#0f172a")
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels([lb[-5:] if len(lb) >= 10 else lb for lb in labels],
                           rotation=35, ha="right", fontsize=9)
        ax.tick_params(axis="y", labelsize=9)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_color("#94a3b8")
        ax.spines["bottom"].set_color("#94a3b8")
        ax.grid(axis="y", linestyle="--", alpha=0.35, color="#94a3b8")
        # Add headroom so data labels never clip
        ymax = max(values + targets) if targets else max(values or [0])
        if ymax > 0:
            ax.set_ylim(0, ymax * 1.14)
        for bar in bars:
            hgt = bar.get_height()
            if hgt > 0:
                ax.text(bar.get_x() + bar.get_width() / 2, hgt, f"{hgt:,.0f}",
                        ha="center", va="bottom", fontsize=8.5, color="#0f172a",
                        weight="semibold")
        plt.tight_layout()

        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        logger.exception("energy chart render failed")
        return None


def _render_losses_chart_png(losses: List[Dict[str, Any]]) -> Optional[bytes]:
    if not _CAPS["charts"] or not losses:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        names = [c.get("label", c.get("id", "?")) for c in losses]
        values = [float(c.get("loss_mwh") or 0) for c in losses]
        palette = ["#ef4444", "#f59e0b", "#3b82f6", "#8b5cf6", "#14b8a6", "#ec4899", "#64748b", "#10b981"]
        colors = [palette[i % len(palette)] for i in range(len(values))]

        fig, ax = plt.subplots(figsize=(9.5, 4.2), dpi=180)
        fig.patch.set_facecolor("white")
        bars = ax.barh(range(len(values)), values, color=colors,
                       edgecolor="white", linewidth=0.8, height=0.68)
        ax.set_yticks(range(len(names)))
        ax.set_yticklabels(names, fontsize=10)
        ax.invert_yaxis()
        ax.set_xlabel("Energy Loss (MWh)", fontsize=11, weight="bold", color="#1e293b")
        ax.set_title("Loss Breakdown by Category", fontsize=13, weight="bold",
                     pad=12, color="#0f172a")
        ax.tick_params(axis="x", labelsize=9)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_color("#94a3b8")
        ax.spines["bottom"].set_color("#94a3b8")
        ax.grid(axis="x", linestyle="--", alpha=0.35, color="#94a3b8")
        vmax = max(values) if values else 0
        if vmax > 0:
            ax.set_xlim(0, vmax * 1.18)
        for bar, v in zip(bars, values):
            if v > 0:
                ax.text(bar.get_width() + vmax * 0.01,
                        bar.get_y() + bar.get_height() / 2,
                        f"{v:,.3f} MWh", va="center", fontsize=9,
                        color="#0f172a", weight="semibold")
        plt.tight_layout()

        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        logger.exception("losses chart render failed")
        return None


def _render_inverter_chart_png(inverters: List[Dict[str, Any]]) -> Optional[bytes]:
    if not _CAPS["charts"] or not inverters:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        # Pick up to the first 30 inverters so the chart stays readable on a single row.
        rows = sorted(inverters, key=lambda r: (r.get("inverter_id") or ""))[:30]
        names = [str(r.get("inverter_id") or "") for r in rows]
        gen = [float(r.get("generation_kwh") or 0) / 1000.0 for r in rows]  # MWh
        # Dashboard bundle emits `pr_pct`; keep legacy name as fallback.
        pr = [float(r.get("pr_pct") or r.get("performance_ratio") or 0) for r in rows]

        fig, ax1 = plt.subplots(figsize=(10.5, 4.2), dpi=180)
        fig.patch.set_facecolor("white")
        idx = list(range(len(names)))
        bars = ax1.bar(idx, gen, color="#0ea5e9", edgecolor="#075985",
                       linewidth=0.6, width=0.68, label="Generation (MWh)")
        ax1.set_xticks(idx)
        ax1.set_xticklabels(names, rotation=55, ha="right", fontsize=8.5)
        ax1.set_ylabel("Generation (MWh)", color="#0369a1", fontsize=11, weight="bold")
        ax1.tick_params(axis="y", colors="#0369a1", labelsize=9)
        ax1.set_title("Per-Inverter Generation & Performance Ratio",
                      fontsize=13, weight="bold", pad=12, color="#0f172a")
        ax1.spines["top"].set_visible(False)
        ax1.spines["left"].set_color("#94a3b8")
        ax1.spines["bottom"].set_color("#94a3b8")
        if gen and max(gen) > 0:
            ax1.set_ylim(0, max(gen) * 1.18)
        # Data value labels above the bars
        for bar, v in zip(bars, gen):
            if v > 0:
                ax1.text(bar.get_x() + bar.get_width() / 2, v,
                         f"{v:,.2f}", ha="center", va="bottom",
                         fontsize=7.5, color="#0c4a6e", weight="semibold", rotation=0)

        ax2 = ax1.twinx()
        ax2.plot(idx, pr, color="#f97316", marker="o", linewidth=2.2,
                 markersize=5, markeredgecolor="white", markeredgewidth=1.0,
                 label="PR (%)", zorder=5)
        ax2.set_ylabel("Performance Ratio (%)", color="#c2410c",
                       fontsize=11, weight="bold")
        ax2.tick_params(axis="y", colors="#c2410c", labelsize=9)
        ax2.set_ylim(0, max(100, max(pr) + 8) if pr else 100)
        ax2.spines["top"].set_visible(False)
        # PR numeric labels near the markers
        for xi, pv in zip(idx, pr):
            if pv > 0:
                ax2.annotate(f"{pv:.1f}", (xi, pv), textcoords="offset points",
                             xytext=(0, 8), ha="center", fontsize=7,
                             color="#9a3412", weight="semibold")

        ax1.grid(axis="y", linestyle="--", alpha=0.3, color="#94a3b8")
        # Combined legend for clarity
        h1, l1 = ax1.get_legend_handles_labels()
        h2, l2 = ax2.get_legend_handles_labels()
        ax1.legend(h1 + h2, l1 + l2, loc="upper right", fontsize=9,
                   frameon=True, facecolor="white", edgecolor="#cbd5e1")
        _ = [bar for bar in bars]
        plt.tight_layout()

        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=140, bbox_inches="tight")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        logger.exception("inverter chart render failed")
        return None


# ──────────────────────────────────────────────────────────────────────────────
# Helper: consistent number formatting
# ──────────────────────────────────────────────────────────────────────────────
def _fmt(v: Any, decimals: int = 2, suffix: str = "") -> str:
    if v is None or v == "":
        return "—"
    try:
        n = float(v)
        if decimals == 0:
            return f"{n:,.0f}{suffix}"
        return f"{n:,.{decimals}f}{suffix}"
    except Exception:
        return str(v)


def _report_title(station: Dict[str, Any], date_from: str, date_to: str, override: Optional[str]) -> str:
    if override:
        return override
    name = station.get("name") or station.get("plant_id") or "Plant"
    return f"Performance Report — {name} — {date_from} to {date_to}"


# ──────────────────────────────────────────────────────────────────────────────
# PDF builder (reportlab)
# ──────────────────────────────────────────────────────────────────────────────
def _build_pdf(payload: Dict[str, Any], req: ReportRequest) -> bytes:
    if not _CAPS["pdf"]:
        raise HTTPException(503, "PDF generation requires the 'reportlab' package. Install it on the server (pip install reportlab).")

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak,
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    bundle = payload["bundle"]
    losses = payload["losses"]
    station = bundle.get("station", {})
    kpis = bundle.get("kpis", {})
    inverters = bundle.get("inverter_performance", [])
    energy = bundle.get("energy", [])
    loss_cats = losses.get("categories", []) if losses else []

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=1.6 * cm, bottomMargin=1.6 * cm,
        title=_report_title(station, req.date_from or "", req.date_to or "", req.title),
        author="Solar Analytics Platform",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Heading1"], fontName="Helvetica-Bold",
                                  fontSize=18, textColor=colors.HexColor("#0f172a"), leading=22, spaceAfter=6)
    subtitle_style = ParagraphStyle("subtitle", parent=styles["BodyText"], fontName="Helvetica",
                                     fontSize=10, textColor=colors.HexColor("#475569"), spaceAfter=14)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName="Helvetica-Bold",
                         fontSize=13, textColor=colors.HexColor("#1e293b"), spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=9.5, leading=13, textColor=colors.HexColor("#334155"))
    small = ParagraphStyle("small", parent=body, fontSize=8, textColor=colors.HexColor("#64748b"))

    story: List[Any] = []

    # Header strip — orange accent bar
    header_table = Table([[Paragraph("<b>SOLAR ANALYTICS • PLANT REPORT</b>", small)]],
                          colWidths=[doc.width])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f97316")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph(_report_title(station, req.date_from or "", req.date_to or "", req.title), title_style))
    subtitle_text = (
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}  •  "
        f"Range {req.date_from or '—'} → {req.date_to or '—'}  •  "
        f"Plant ID: {station.get('plant_id', req.plant_id)}"
    )
    story.append(Paragraph(subtitle_text, subtitle_style))

    # ── Overview ──
    if "overview" in req.sections:
        story.append(Paragraph("Plant Overview", h2))
        overview_rows = [
            ["Name", station.get("name") or "—", "Technology", station.get("technology") or "—"],
            ["Capacity (MWp)", _fmt(station.get("capacity_mwp"), 3), "Status", station.get("status") or "—"],
            ["Location", station.get("location") or "—", "COD", station.get("cod_date") or "—"],
            ["Plant Age (yrs)", _fmt(station.get("plant_age_years"), 1), "PPA Tariff", _fmt(station.get("ppa_tariff"), 2)],
        ]
        t = Table(overview_rows, colWidths=[3.2 * cm, 5.5 * cm, 3.2 * cm, 5.5 * cm], hAlign="LEFT")
        t.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f1f5f9")),
            ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#f1f5f9")),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1e293b")),
            ("TEXTCOLOR", (2, 0), (2, -1), colors.HexColor("#1e293b")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(t)

    # ── KPI tiles ──
    if "kpis" in req.sections:
        story.append(Paragraph("Key Performance Indicators", h2))
        kpi_items = [
            ("Total Generation",       _fmt(kpis.get("total_inverter_generation_mwh"), 2), "MWh",  "#10b981"),
            ("Energy Export",          _fmt(kpis.get("energy_export_mwh"), 2),              "MWh",  "#3b82f6"),
            ("Peak Power",             _fmt(kpis.get("peak_power_kw"), 1),                  "kW",   "#f59e0b"),
            ("Avg Active Power",       _fmt(kpis.get("active_power_kw"), 1),                "kW",   "#06b6d4"),
            ("Performance Ratio",      _fmt(kpis.get("performance_ratio"), 2),              "%",    "#8b5cf6"),
            ("Plant Load Factor",      _fmt(kpis.get("plant_load_factor"), 2),              "%",    "#ec4899"),
            ("Insolation (GTI)",       _fmt(kpis.get("insolation_kwh_m2"), 2),              "kWh/m²","#f97316"),
        ]
        # Render as a 4-column grid of tiles
        rows_per: List[List[Any]] = []
        row: List[Any] = []
        for label, value, unit, color in kpi_items:
            cell = Table([
                [Paragraph(f"<font size=7 color='#64748b'><b>{label.upper()}</b></font>", body)],
                [Paragraph(f"<font size=18 color='{color}'><b>{value}</b></font>", body)],
                [Paragraph(f"<font size=8 color='#64748b'>{unit}</font>", body)],
            ], colWidths=[4.2 * cm])
            cell.setStyle(TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LINEBEFORE", (0, 0), (0, -1), 2, colors.HexColor(color)),
            ]))
            row.append(cell)
            if len(row) == 4:
                rows_per.append(row)
                row = []
        if row:
            while len(row) < 4:
                row.append("")
            rows_per.append(row)
        grid = Table(rows_per, colWidths=[4.2 * cm] * 4, hAlign="LEFT")
        grid.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(grid)

    # ── Energy trend ──
    if "energy_trend" in req.sections and energy:
        story.append(Paragraph("Daily Energy Generation", h2))
        png = _render_energy_chart_png(energy)
        if png:
            story.append(RLImage(BytesIO(png), width=doc.width, height=7.5 * cm))
        else:
            story.append(Paragraph("Chart rendering unavailable (matplotlib missing on server).", small))

    # ── Losses ──
    if "losses" in req.sections and loss_cats:
        story.append(PageBreak())
        story.append(Paragraph("Loss Analysis by Category", h2))
        png = _render_losses_chart_png(loss_cats)
        if png:
            story.append(RLImage(BytesIO(png), width=doc.width, height=8 * cm))
        rows = [["Category", "Loss (MWh)", "Fault Count", "Note"]]
        for c in loss_cats:
            rows.append([
                c.get("label", c.get("id", "—")),
                _fmt(c.get("loss_mwh"), 3),
                _fmt(c.get("fault_count"), 0),
                Paragraph(f"<font size=8 color='#64748b'>{c.get('metric_note') or ''}</font>", body),
            ])
        totals = losses.get("totals", {}) if losses else {}
        rows.append(["TOTAL", _fmt(totals.get("loss_mwh"), 3), _fmt(totals.get("fault_count"), 0), ""])
        t = Table(rows, colWidths=[4.0 * cm, 2.5 * cm, 2.3 * cm, 8.0 * cm], repeatRows=1)
        t.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#fee2e2")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("ALIGN", (1, 1), (2, -1), "RIGHT"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f8fafc")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(Spacer(1, 8))
        story.append(t)

    # ── Inverter performance ──
    if "inverters" in req.sections and inverters:
        story.append(PageBreak())
        story.append(Paragraph("Inverter Performance", h2))
        png = _render_inverter_chart_png(inverters)
        if png:
            story.append(RLImage(BytesIO(png), width=doc.width, height=7.5 * cm))
        rows = [["Inverter", "Gen (kWh)", "DC Cap (kWp)", "Yield (kWh/kWp)", "PR (%)", "PLF (%)", "η (%)"]]
        for r in inverters:
            rows.append([
                r.get("inverter_id") or "—",
                _fmt(r.get("generation_kwh"), 1),
                _fmt(r.get("dc_capacity_kwp"), 2),
                _fmt(r.get("yield_kwh_kwp"), 2),
                _fmt(r.get("pr_pct"), 2),
                _fmt(r.get("plf_pct"), 2),
                _fmt(r.get("efficiency_pct"), 1),
            ])
        t = Table(rows, colWidths=[3.0 * cm, 2.4 * cm, 2.6 * cm, 2.6 * cm, 2.0 * cm, 2.0 * cm, 1.8 * cm], repeatRows=1)
        t.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(Spacer(1, 8))
        story.append(t)

    # Footer on every page
    def _on_page(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#94a3b8"))
        canvas.setFont("Helvetica", 8)
        footer = f"Solar Analytics Platform  •  Page {canvas.getPageNumber()}  •  {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        canvas.drawRightString(A4[0] - 1.8 * cm, 1.0 * cm, footer)
        canvas.restoreState()

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# XLSX builder (openpyxl)
# ──────────────────────────────────────────────────────────────────────────────
def _build_xlsx(payload: Dict[str, Any], req: ReportRequest) -> bytes:
    if not _CAPS["xlsx"]:
        raise HTTPException(503, "XLSX generation requires 'openpyxl' (already listed in requirements).")

    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    bundle = payload["bundle"]
    losses = payload["losses"]
    station = bundle.get("station", {})
    kpis = bundle.get("kpis", {})
    inverters = bundle.get("inverter_performance", [])
    energy = bundle.get("energy", [])
    loss_cats = losses.get("categories", []) if losses else []

    wb = Workbook()

    # Shared styles
    title_font = Font(name="Calibri", size=16, bold=True, color="FFFFFF")
    head_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    label_font = Font(name="Calibri", size=10, bold=True, color="1E293B")
    value_font = Font(name="Calibri", size=11, color="0F172A")
    header_fill = PatternFill("solid", fgColor="F97316")
    table_head_fill = PatternFill("solid", fgColor="1E293B")
    band_fill = PatternFill("solid", fgColor="F8FAFC")
    thin = Side(border_style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _apply_column_widths(ws, widths):
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    # ── Sheet 1: Overview ──
    ws = wb.active
    ws.title = "Overview"
    ws.merge_cells("A1:D1")
    ws["A1"] = _report_title(station, req.date_from or "", req.date_to or "", req.title)
    ws["A1"].font = title_font
    ws["A1"].fill = header_fill
    ws["A1"].alignment = Alignment(horizontal="left", vertical="center", indent=1)
    ws.row_dimensions[1].height = 28

    ws["A2"] = f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}   •   Range {req.date_from or '—'} → {req.date_to or '—'}"
    ws.merge_cells("A2:D2")
    ws["A2"].font = Font(size=10, color="64748B")
    ws.row_dimensions[2].height = 18

    overview = [
        ("Plant ID", station.get("plant_id", req.plant_id), "Technology", station.get("technology") or "—"),
        ("Name", station.get("name") or "—", "Status", station.get("status") or "—"),
        ("Capacity (MWp)", station.get("capacity_mwp"), "Location", station.get("location") or "—"),
        ("COD", station.get("cod_date") or "—", "Plant Age (yrs)", station.get("plant_age_years")),
        ("PPA Tariff", station.get("ppa_tariff"), "", ""),
    ]
    for i, (k1, v1, k2, v2) in enumerate(overview, start=4):
        ws.cell(row=i, column=1, value=k1).font = label_font
        ws.cell(row=i, column=2, value=v1).font = value_font
        ws.cell(row=i, column=3, value=k2).font = label_font
        ws.cell(row=i, column=4, value=v2).font = value_font
        for c in range(1, 5):
            ws.cell(row=i, column=c).border = border

    # KPIs block
    start = 4 + len(overview) + 1
    ws.cell(row=start, column=1, value="KEY PERFORMANCE INDICATORS").font = Font(bold=True, size=11, color="1E293B")
    start += 1
    kpi_items = [
        ("Total Generation (MWh)", kpis.get("total_inverter_generation_mwh")),
        ("Energy Export (MWh)", kpis.get("energy_export_mwh")),
        ("Peak Power (kW)", kpis.get("peak_power_kw")),
        ("Avg Active Power (kW)", kpis.get("active_power_kw")),
        ("Performance Ratio (%)", kpis.get("performance_ratio")),
        ("Plant Load Factor (%)", kpis.get("plant_load_factor")),
        ("Insolation (kWh/m²)", kpis.get("insolation_kwh_m2")),
    ]
    for i, (k, v) in enumerate(kpi_items):
        r = start + i
        ws.cell(row=r, column=1, value=k).font = label_font
        cell = ws.cell(row=r, column=2, value=v)
        cell.font = Font(size=12, bold=True, color="1E40AF")
        cell.number_format = "#,##0.00"
        for c in range(1, 5):
            ws.cell(row=r, column=c).border = border
            if (i % 2) == 1:
                ws.cell(row=r, column=c).fill = band_fill

    _apply_column_widths(ws, [26, 22, 22, 28])

    # ── Sheet 2: Inverters ──
    if inverters:
        ws2 = wb.create_sheet("Inverters")
        headers = ["Inverter", "Generation (kWh)", "DC Capacity (kWp)", "Yield (kWh/kWp)", "PR (%)", "PLF (%)", "Efficiency (%)"]
        for i, hdr in enumerate(headers, 1):
            cell = ws2.cell(row=1, column=i, value=hdr)
            cell.font = head_font
            cell.fill = table_head_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
        ws2.row_dimensions[1].height = 22
        for idx, r in enumerate(inverters, start=2):
            vals = [
                r.get("inverter_id") or "—",
                r.get("generation_kwh"),
                r.get("dc_capacity_kwp"),
                r.get("yield_kwh_kwp"),
                r.get("pr_pct"),
                r.get("plf_pct"),
                r.get("efficiency_pct"),
            ]
            for c, val in enumerate(vals, 1):
                cell = ws2.cell(row=idx, column=c, value=val)
                cell.border = border
                cell.font = Font(size=10)
                if c > 1:
                    cell.alignment = Alignment(horizontal="right")
                    cell.number_format = "#,##0.00"
                if (idx % 2) == 0:
                    cell.fill = band_fill
        _apply_column_widths(ws2, [16, 18, 18, 18, 14, 14, 16])
        ws2.freeze_panes = "A2"

    # ── Sheet 3: Losses ──
    if loss_cats:
        ws3 = wb.create_sheet("Losses")
        headers = ["Category", "Loss (MWh)", "Fault Count", "Note"]
        for i, hdr in enumerate(headers, 1):
            cell = ws3.cell(row=1, column=i, value=hdr)
            cell.font = head_font
            cell.fill = table_head_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
        ws3.row_dimensions[1].height = 22
        for idx, c in enumerate(loss_cats, start=2):
            vals = [c.get("label", c.get("id", "—")), c.get("loss_mwh"), c.get("fault_count"), c.get("metric_note") or ""]
            for j, val in enumerate(vals, 1):
                cell = ws3.cell(row=idx, column=j, value=val)
                cell.border = border
                cell.font = Font(size=10)
                if j in (2, 3):
                    cell.alignment = Alignment(horizontal="right")
                    cell.number_format = "#,##0.000" if j == 2 else "#,##0"
                if (idx % 2) == 0:
                    cell.fill = band_fill
        # total row
        totals = losses.get("totals", {}) if losses else {}
        r = 2 + len(loss_cats)
        ws3.cell(row=r, column=1, value="TOTAL").font = Font(size=11, bold=True)
        ws3.cell(row=r, column=2, value=totals.get("loss_mwh")).number_format = "#,##0.000"
        ws3.cell(row=r, column=2).font = Font(size=11, bold=True)
        ws3.cell(row=r, column=3, value=totals.get("fault_count")).font = Font(size=11, bold=True)
        for c in range(1, 5):
            cell = ws3.cell(row=r, column=c)
            cell.fill = PatternFill("solid", fgColor="FEE2E2")
            cell.border = border
        _apply_column_widths(ws3, [24, 14, 14, 60])
        ws3.freeze_panes = "A2"

    # ── Sheet 4: Energy Trend ──
    if energy:
        ws4 = wb.create_sheet("Energy Trend")
        headers = ["Date", "Energy (kWh)", "Target (kWh)"]
        for i, hdr in enumerate(headers, 1):
            cell = ws4.cell(row=1, column=i, value=hdr)
            cell.font = head_font
            cell.fill = table_head_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
        for idx, p in enumerate(energy, start=2):
            ws4.cell(row=idx, column=1, value=p.get("date") or p.get("day"))
            cell = ws4.cell(row=idx, column=2, value=p.get("actual_kwh") or p.get("energy_kwh"))
            cell.number_format = "#,##0.00"
            cell3 = ws4.cell(row=idx, column=3, value=p.get("target_kwh") or p.get("target"))
            cell3.number_format = "#,##0.00"
            for c in range(1, 4):
                ws4.cell(row=idx, column=c).border = border
                if (idx % 2) == 0:
                    ws4.cell(row=idx, column=c).fill = band_fill
        _apply_column_widths(ws4, [14, 18, 18])
        ws4.freeze_panes = "A2"

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# DOCX builder (python-docx)
# ──────────────────────────────────────────────────────────────────────────────
def _build_docx(payload: Dict[str, Any], req: ReportRequest) -> bytes:
    if not _CAPS["docx"]:
        raise HTTPException(503, "DOCX generation requires the 'python-docx' package. Install it on the server (pip install python-docx).")

    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    bundle = payload["bundle"]
    losses = payload["losses"]
    station = bundle.get("station", {})
    kpis = bundle.get("kpis", {})
    inverters = bundle.get("inverter_performance", [])
    energy = bundle.get("energy", [])
    loss_cats = losses.get("categories", []) if losses else []

    doc = Document()
    # Page margins
    for sec in doc.sections:
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)
        sec.top_margin = Cm(1.6)
        sec.bottom_margin = Cm(1.6)

    def _shade_cell(cell, hex_color: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    # Title band
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("SOLAR ANALYTICS • PLANT REPORT")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)

    title = doc.add_heading(_report_title(station, req.date_from or "", req.date_to or "", req.title), level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}   •   "
        f"Range {req.date_from or '—'} → {req.date_to or '—'}   •   "
        f"Plant ID: {station.get('plant_id', req.plant_id)}"
    ).font.size = Pt(9)

    # ── Overview table ──
    if "overview" in req.sections:
        doc.add_heading("Plant Overview", level=1)
        overview_rows = [
            ("Name", station.get("name") or "—", "Technology", station.get("technology") or "—"),
            ("Capacity (MWp)", _fmt(station.get("capacity_mwp"), 3), "Status", station.get("status") or "—"),
            ("Location", station.get("location") or "—", "COD", station.get("cod_date") or "—"),
            ("Plant Age (yrs)", _fmt(station.get("plant_age_years"), 1), "PPA Tariff", _fmt(station.get("ppa_tariff"), 2)),
        ]
        tbl = doc.add_table(rows=len(overview_rows), cols=4)
        tbl.style = "Light Grid Accent 1"
        for i, (k1, v1, k2, v2) in enumerate(overview_rows):
            cells = tbl.rows[i].cells
            cells[0].text = str(k1); _shade_cell(cells[0], "F1F5F9")
            cells[1].text = str(v1)
            cells[2].text = str(k2); _shade_cell(cells[2], "F1F5F9")
            cells[3].text = str(v2)
            for idx in (0, 2):
                for r in cells[idx].paragraphs[0].runs:
                    r.bold = True

    # ── KPI panel ──
    if "kpis" in req.sections:
        doc.add_heading("Key Performance Indicators", level=1)
        kpi_items = [
            ("Total Generation (MWh)", _fmt(kpis.get("total_inverter_generation_mwh"), 2), "10B981"),
            ("Energy Export (MWh)",    _fmt(kpis.get("energy_export_mwh"), 2),              "3B82F6"),
            ("Peak Power (kW)",         _fmt(kpis.get("peak_power_kw"), 1),                 "F59E0B"),
            ("Avg Active Power (kW)",   _fmt(kpis.get("active_power_kw"), 1),               "06B6D4"),
            ("Performance Ratio (%)",   _fmt(kpis.get("performance_ratio"), 2),             "8B5CF6"),
            ("Plant Load Factor (%)",   _fmt(kpis.get("plant_load_factor"), 2),             "EC4899"),
            ("Insolation (kWh/m²)",     _fmt(kpis.get("insolation_kwh_m2"), 2),             "F97316"),
        ]
        cols = 4
        rows = (len(kpi_items) + cols - 1) // cols
        kt = doc.add_table(rows=rows, cols=cols)
        kt.autofit = True
        for i, (label, value, color) in enumerate(kpi_items):
            r, c = i // cols, i % cols
            cell = kt.rows[r].cells[c]
            p1 = cell.paragraphs[0]
            run_label = p1.add_run(label.upper() + "\n")
            run_label.font.size = Pt(7)
            run_label.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            run_label.bold = True
            run_val = p1.add_run(value)
            run_val.bold = True
            run_val.font.size = Pt(16)
            run_val.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
            _shade_cell(cell, "FFFFFF")

    # ── Energy trend chart ──
    if "energy_trend" in req.sections and energy:
        doc.add_heading("Daily Energy Generation", level=1)
        png = _render_energy_chart_png(energy)
        if png:
            doc.add_picture(BytesIO(png), width=Inches(6.5))
        else:
            doc.add_paragraph("(matplotlib missing — chart skipped)").runs[0].font.size = Pt(9)

    # ── Losses ──
    if "losses" in req.sections and loss_cats:
        doc.add_heading("Loss Analysis by Category", level=1)
        png = _render_losses_chart_png(loss_cats)
        if png:
            doc.add_picture(BytesIO(png), width=Inches(6.5))
        tbl = doc.add_table(rows=1 + len(loss_cats) + 1, cols=4)
        tbl.style = "Light Grid Accent 1"
        head = tbl.rows[0].cells
        for i, hdr in enumerate(["Category", "Loss (MWh)", "Fault Count", "Note"]):
            head[i].text = hdr
            _shade_cell(head[i], "1E293B")
            for r in head[i].paragraphs[0].runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, c in enumerate(loss_cats, start=1):
            row = tbl.rows[i].cells
            row[0].text = str(c.get("label", c.get("id", "—")))
            row[1].text = _fmt(c.get("loss_mwh"), 3)
            row[2].text = _fmt(c.get("fault_count"), 0)
            row[3].text = str(c.get("metric_note") or "")
        totals = losses.get("totals", {}) if losses else {}
        last = tbl.rows[-1].cells
        last[0].text = "TOTAL"
        last[1].text = _fmt(totals.get("loss_mwh"), 3)
        last[2].text = _fmt(totals.get("fault_count"), 0)
        for cell in last:
            _shade_cell(cell, "FEE2E2")
            for r in cell.paragraphs[0].runs:
                r.bold = True

    # ── Inverter performance ──
    if "inverters" in req.sections and inverters:
        doc.add_heading("Inverter Performance", level=1)
        png = _render_inverter_chart_png(inverters)
        if png:
            doc.add_picture(BytesIO(png), width=Inches(6.5))

        headers = ["Inverter", "Gen (kWh)", "DC Cap (kWp)", "Yield (kWh/kWp)", "PR (%)", "PLF (%)", "η (%)"]
        tbl = doc.add_table(rows=1 + len(inverters), cols=len(headers))
        tbl.style = "Light Grid Accent 1"
        for i, hdr in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = hdr
            _shade_cell(cell, "1E293B")
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, r in enumerate(inverters, start=1):
            row = tbl.rows[i].cells
            row[0].text = str(r.get("inverter_id") or "—")
            row[1].text = _fmt(r.get("generation_kwh"), 1)
            row[2].text = _fmt(r.get("dc_capacity_kwp"), 2)
            row[3].text = _fmt(r.get("yield_kwh_kwp"), 2)
            row[4].text = _fmt(r.get("pr_pct"), 2)
            row[5].text = _fmt(r.get("plf_pct"), 2)
            row[6].text = _fmt(r.get("efficiency_pct"), 1)

    # Footer
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"Solar Analytics Platform   •   {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    for r in footer.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# HTML builder — self-contained, printable, base64-embedded charts
# ──────────────────────────────────────────────────────────────────────────────
def _html_escape(s: Any) -> str:
    """Minimal HTML escape with None-safe behaviour."""
    import html as _html

    return _html.escape("" if s is None else str(s), quote=True)


def _png_to_data_uri(png: Optional[bytes]) -> str:
    if not png:
        return ""
    return "data:image/png;base64," + base64.b64encode(png).decode("ascii")


def _build_html(payload: Dict[str, Any], req: ReportRequest) -> bytes:
    """Produce a standalone, printable HTML report (single file, no external assets)."""
    bundle = payload["bundle"]
    losses = payload["losses"]
    station = bundle.get("station", {})
    kpis = bundle.get("kpis", {})
    inverters = bundle.get("inverter_performance", [])
    energy = bundle.get("energy", [])
    loss_cats = losses.get("categories", []) if losses else []
    totals = losses.get("totals", {}) if losses else {}

    e = _html_escape
    title = _report_title(station, req.date_from or "", req.date_to or "", req.title)
    gen_ts = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Chart data URIs (empty string if matplotlib missing — the HTML gracefully
    # skips the <img> instead of showing a broken-image icon).
    energy_uri = _png_to_data_uri(_render_energy_chart_png(energy)) if "energy_trend" in req.sections else ""
    losses_uri = _png_to_data_uri(_render_losses_chart_png(loss_cats)) if "losses" in req.sections else ""
    inv_uri = _png_to_data_uri(_render_inverter_chart_png(inverters)) if "inverters" in req.sections else ""

    # ── Helpers for table rows ──
    def _row(cols: List[Any]) -> str:
        return "<tr>" + "".join(f"<td>{c}</td>" for c in cols) + "</tr>"

    def _kpi_tile(label: str, value: str, unit: str, color: str) -> str:
        return (
            f'<div class="kpi-tile" style="--accent:{color}">'
            f'<div class="kpi-label">{e(label)}</div>'
            f'<div class="kpi-value">{e(value)}</div>'
            f'<div class="kpi-unit">{e(unit)}</div>'
            "</div>"
        )

    # Overview block
    overview_html = ""
    if "overview" in req.sections:
        ov_rows = [
            ("Name", station.get("name") or "—", "Technology", station.get("technology") or "—"),
            ("Capacity (MWp)", _fmt(station.get("capacity_mwp"), 3), "Status", station.get("status") or "—"),
            ("Location", station.get("location") or "—", "COD", station.get("cod_date") or "—"),
            ("Plant Age (yrs)", _fmt(station.get("plant_age_years"), 1), "PPA Tariff", _fmt(station.get("ppa_tariff"), 2)),
        ]
        ov_trs = "".join(
            f"<tr><th>{e(k1)}</th><td>{e(v1)}</td><th>{e(k2)}</th><td>{e(v2)}</td></tr>"
            for k1, v1, k2, v2 in ov_rows
        )
        overview_html = (
            '<section class="card"><h2>Plant Overview</h2>'
            f'<table class="kv">{ov_trs}</table></section>'
        )

    # KPI block
    kpi_html = ""
    if "kpis" in req.sections:
        kpi_items = [
            ("Total Generation", _fmt(kpis.get("total_inverter_generation_mwh"), 2), "MWh",    "#10b981"),
            ("Energy Export",    _fmt(kpis.get("energy_export_mwh"), 2),              "MWh",    "#3b82f6"),
            ("Peak Power",       _fmt(kpis.get("peak_power_kw"), 1),                  "kW",     "#f59e0b"),
            ("Avg Active Power", _fmt(kpis.get("active_power_kw"), 1),                "kW",     "#06b6d4"),
            ("Performance Ratio", _fmt(kpis.get("performance_ratio"), 2),             "%",      "#8b5cf6"),
            ("Plant Load Factor", _fmt(kpis.get("plant_load_factor"), 2),             "%",      "#ec4899"),
            ("Insolation (GTI)", _fmt(kpis.get("insolation_kwh_m2"), 2),              "kWh/m²", "#f97316"),
        ]
        tiles = "".join(_kpi_tile(*it) for it in kpi_items)
        kpi_html = (
            '<section class="card"><h2>Key Performance Indicators</h2>'
            f'<div class="kpi-grid">{tiles}</div></section>'
        )

    # Energy trend block
    energy_html = ""
    if "energy_trend" in req.sections and energy:
        if energy_uri:
            img = f'<img class="chart" src="{energy_uri}" alt="Daily Energy Generation chart"/>'
        else:
            img = '<p class="muted">Chart rendering unavailable (matplotlib missing on server).</p>'
        rows_html = "".join(
            _row([
                e(p.get("date") or p.get("day") or ""),
                _fmt(p.get("actual_kwh") or p.get("energy_kwh"), 1),
                _fmt(p.get("target_kwh") or p.get("target"), 1),
            ])
            for p in energy
        )
        energy_html = (
            '<section class="card"><h2>Daily Energy Generation</h2>'
            f"{img}"
            '<table class="data"><thead><tr>'
            "<th>Date</th><th>Actual (kWh)</th><th>Target (kWh)</th>"
            f"</tr></thead><tbody>{rows_html}</tbody></table></section>"
        )

    # Losses block
    losses_html = ""
    if "losses" in req.sections and loss_cats:
        chart = f'<img class="chart" src="{losses_uri}" alt="Loss Breakdown chart"/>' if losses_uri else ""
        body_rows = "".join(
            _row([
                e(c.get("label", c.get("id", "—"))),
                f'<span class="num">{_fmt(c.get("loss_mwh"), 3)}</span>',
                f'<span class="num">{_fmt(c.get("fault_count"), 0)}</span>',
                f'<span class="muted small">{e(c.get("metric_note") or "")}</span>',
            ])
            for c in loss_cats
        )
        totals_row = _row([
            "<strong>TOTAL</strong>",
            f'<span class="num"><strong>{_fmt(totals.get("loss_mwh"), 3)}</strong></span>',
            f'<span class="num"><strong>{_fmt(totals.get("fault_count"), 0)}</strong></span>',
            "",
        ])
        losses_html = (
            '<section class="card"><h2>Loss Analysis by Category</h2>'
            f"{chart}"
            '<table class="data losses"><thead><tr>'
            "<th>Category</th><th>Loss (MWh)</th><th>Fault Count</th><th>Note</th>"
            f"</tr></thead><tbody>{body_rows}<tr class='total'>{totals_row[4:]}</tbody></table></section>"
        )
        # NOTE: the slice trick above drops the leading '<tr>' from totals_row so we can
        # re-open <tr> with a class.

    # Inverter block
    inv_html = ""
    if "inverters" in req.sections and inverters:
        chart = f'<img class="chart" src="{inv_uri}" alt="Per-inverter chart"/>' if inv_uri else ""
        inv_rows = "".join(
            _row([
                e(r.get("inverter_id") or "—"),
                f'<span class="num">{_fmt(r.get("generation_kwh"), 1)}</span>',
                f'<span class="num">{_fmt(r.get("dc_capacity_kwp"), 2)}</span>',
                f'<span class="num">{_fmt(r.get("yield_kwh_kwp"), 2)}</span>',
                f'<span class="num">{_fmt(r.get("pr_pct"), 2)}</span>',
                f'<span class="num">{_fmt(r.get("plf_pct"), 2)}</span>',
                f'<span class="num">{_fmt(r.get("efficiency_pct"), 1)}</span>',
            ])
            for r in inverters
        )
        inv_html = (
            '<section class="card"><h2>Inverter Performance</h2>'
            f"{chart}"
            '<table class="data inv"><thead><tr>'
            "<th>Inverter</th><th>Gen (kWh)</th><th>DC Cap (kWp)</th><th>Yield (kWh/kWp)</th>"
            "<th>PR (%)</th><th>PLF (%)</th><th>η (%)</th>"
            f"</tr></thead><tbody>{inv_rows}</tbody></table></section>"
        )

    # Styles — a compact, print-friendly modern look that also works beautifully
    # on screens. All colours and spacing mirror the on-screen platform theme.
    style = """
    :root{
      --bg:#f5f7fb;--surface:#ffffff;--ink:#0f172a;--ink-2:#334155;--muted:#64748b;
      --accent:#f97316;--primary:#1e293b;--ring:#e2e8f0;--band:#f8fafc;--shadow:0 1px 2px rgba(15,23,42,.06),0 4px 14px rgba(15,23,42,.04);
    }
    *{box-sizing:border-box}
    html,body{margin:0;padding:0;background:var(--bg);color:var(--ink);font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;-webkit-font-smoothing:antialiased}
    .wrap{max-width:1180px;margin:0 auto;padding:28px 32px 48px}
    .banner{background:var(--accent);color:#fff;font-weight:700;letter-spacing:.08em;font-size:11px;padding:8px 14px;border-radius:8px;display:inline-block;margin-bottom:14px}
    h1.title{font-size:28px;margin:0 0 6px;font-weight:800;color:var(--ink);letter-spacing:-.01em}
    .subtitle{color:var(--muted);font-size:13px;margin:0 0 22px}
    .card{background:var(--surface);border:1px solid var(--ring);border-radius:14px;padding:22px 24px;margin:14px 0;box-shadow:var(--shadow)}
    h2{font-size:16px;margin:0 0 14px;color:var(--primary);font-weight:700;letter-spacing:.01em}
    .kv{border-collapse:collapse;width:100%}
    .kv th,.kv td{padding:9px 12px;border:1px solid var(--ring);font-size:13px;text-align:left}
    .kv th{background:#f1f5f9;color:var(--ink);font-weight:600;width:18%}
    .kpi-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:12px}
    .kpi-tile{border:1px solid var(--ring);border-left:4px solid var(--accent,#f97316);border-radius:10px;padding:14px 16px;background:#fff}
    .kpi-label{font-size:10px;letter-spacing:.12em;font-weight:700;color:var(--muted);text-transform:uppercase}
    .kpi-value{font-size:26px;font-weight:800;color:var(--accent,#1e293b);margin-top:2px;line-height:1.1}
    .kpi-unit{font-size:11px;color:var(--muted);margin-top:1px}
    img.chart{width:100%;height:auto;max-height:420px;object-fit:contain;border:1px solid var(--ring);border-radius:10px;background:#fff;padding:6px;margin-bottom:12px}
    table.data{border-collapse:collapse;width:100%;font-size:12.5px}
    table.data th,table.data td{border:1px solid var(--ring);padding:7px 10px;text-align:left}
    table.data th{background:var(--primary);color:#fff;font-weight:600;font-size:12px;letter-spacing:.02em}
    table.data tbody tr:nth-child(even){background:var(--band)}
    table.data tbody tr:hover{background:#eff6ff}
    table.data .num{font-variant-numeric:tabular-nums;text-align:right;display:block}
    table.losses tr.total td{background:#fee2e2;font-weight:700;color:#7f1d1d}
    .small{font-size:11px}
    .muted{color:var(--muted)}
    .footer{margin-top:28px;color:var(--muted);font-size:11px;text-align:right;border-top:1px solid var(--ring);padding-top:10px}
    .toolbar{display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:10px;margin-bottom:14px}
    .btn-print{background:var(--primary);color:#fff;border:0;border-radius:8px;padding:8px 14px;font-size:12px;font-weight:600;cursor:pointer;letter-spacing:.03em}
    .btn-print:hover{background:#0f172a}
    @media print{
      body{background:#fff}.wrap{padding:6mm 4mm}.card{break-inside:avoid;box-shadow:none;border-color:#e2e8f0}
      .toolbar,.btn-print{display:none}img.chart{max-height:none}
      h2{break-after:avoid}section{break-inside:avoid-page}
    }
    @media (max-width:880px){.kpi-grid{grid-template-columns:repeat(2,1fr)}.kv th{width:32%}}
    """

    doc = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>{e(title)}</title>
<style>{style}</style>
</head>
<body>
<div class="wrap">
  <div class="toolbar">
    <span class="banner">SOLAR ANALYTICS • PLANT REPORT</span>
    <button class="btn-print" onclick="window.print()">Print / Save as PDF</button>
  </div>
  <h1 class="title">{e(title)}</h1>
  <p class="subtitle">Generated {e(gen_ts)} &nbsp;•&nbsp; Range {e(req.date_from or "—")} → {e(req.date_to or "—")} &nbsp;•&nbsp; Plant ID: {e(station.get("plant_id") or req.plant_id)}</p>
  {overview_html}
  {kpi_html}
  {energy_html}
  {losses_html}
  {inv_html}
  <div class="footer">Solar Analytics Platform • Generated {e(gen_ts)}</div>
</div>
</body>
</html>"""
    return doc.encode("utf-8")


# ──────────────────────────────────────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────────────────────────────────────
@router.get("/options")
def report_options(current_user: User = Depends(get_current_user)):
    """Expose which output formats this server instance actually supports so the UI
    can hide / disable unavailable options (e.g. reportlab not installed)."""
    return {
        "formats": {
            "pdf": _CAPS["pdf"],
            "xlsx": _CAPS["xlsx"],
            "docx": _CAPS["docx"],
            "html": _CAPS["html"],
        },
        "charts_enabled": _CAPS["charts"],
        "sections": [
            {"id": "overview",     "label": "Plant Overview"},
            {"id": "kpis",         "label": "Key Performance Indicators"},
            {"id": "energy_trend", "label": "Daily Energy Trend"},
            {"id": "inverters",    "label": "Per-Inverter Performance"},
            {"id": "losses",       "label": "Loss Analysis"},
            {"id": "faults",       "label": "Fault Summary"},  # rolls into losses section today
        ],
    }


@router.post("/generate")
def generate_report(
    req: ReportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return a streaming response with the generated document."""
    # Normalize dates
    if not req.date_from or not req.date_to:
        today = date.today()
        req.date_from = req.date_from or today.replace(day=1).isoformat()
        req.date_to = req.date_to or today.isoformat()

    # Validate plant exists (keeps the error message crisp if plant_id is bad)
    if not db.query(Plant).filter(Plant.plant_id == req.plant_id).first():
        raise HTTPException(404, f"Unknown plant_id: {req.plant_id}")

    payload = _gather_report_payload(db, req.plant_id, req.date_from, req.date_to)

    if req.format == "pdf":
        data = _build_pdf(payload, req)
        media = "application/pdf"
        ext = "pdf"
    elif req.format == "xlsx":
        data = _build_xlsx(payload, req)
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ext = "xlsx"
    elif req.format == "docx":
        data = _build_docx(payload, req)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif req.format == "html":
        data = _build_html(payload, req)
        media = "text/html; charset=utf-8"
        ext = "html"
    else:
        raise HTTPException(400, f"Unsupported format: {req.format}")

    plant_slug = (req.plant_id or "plant").replace("/", "_").replace("\\", "_")
    filename = f"report_{plant_slug}_{req.date_from}_{req.date_to}.{ext}"
    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"',
        "X-Report-Format": req.format,
        "X-Report-Bytes": str(len(data)),
    }
    return StreamingResponse(BytesIO(data), media_type=media, headers=headers)
